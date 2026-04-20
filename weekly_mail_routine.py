#!/usr/bin/env python
# -*- coding: utf-8 -*-

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import json
import os
import pickle
import base64
import io
from datetime import datetime, timedelta
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
import googleapiclient.discovery as discovery
from googleapiclient.http import MediaFileUpload, MediaIoBaseUpload
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 설정
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly',
          'https://www.googleapis.com/auth/calendar',
          'https://www.googleapis.com/auth/drive']

# 인증 정보 경로 설정 (환경 변수 또는 기본값)
CREDENTIALS_FILE = os.getenv('GOOGLE_CREDENTIALS_FILE', 'G:\\내 드라이브\\01_토닥\\01_개발\\99_claude\\00_default\\gmail.json')
TOKEN_FILE = os.getenv('TOKEN_FILE', 'token.pickle')
STATE_FILE = os.getenv('STATE_FILE', 'mail_routine_state.json')
MAIL_FOLDER_ID = '1_E3GtIKt1TpPpUa8m1KTW5zqJTE_nQom'

# 환경 변수에서 Google 인증 정보 읽기 (원격 실행 시)
def setup_credentials():
    """환경 변수에서 인증 정보 설정"""
    google_creds_json = os.getenv('GOOGLE_CREDENTIALS')
    if google_creds_json:
        global CREDENTIALS_FILE
        creds_path = 'gmail_credentials.json'
        with open(creds_path, 'w') as f:
            f.write(google_creds_json)
        CREDENTIALS_FILE = creds_path

def get_services():
    """Google API 서비스 인증"""
    creds = None
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'rb') as token:
            creds = pickle.load(token)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_FILE, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(TOKEN_FILE, 'wb') as token:
            pickle.dump(creds, token)

    gmail = discovery.build('gmail', 'v1', credentials=creds)
    calendar = discovery.build('calendar', 'v3', credentials=creds)
    drive = discovery.build('drive', 'v3', credentials=creds)

    return gmail, calendar, drive, creds

def get_last_processed_time():
    """마지막 처리 시간 조회"""
    if os.path.exists(STATE_FILE):
        with open(STATE_FILE, 'r', encoding='utf-8') as f:
            state = json.load(f)
            return state.get('last_processed_time')
    return None

def save_last_processed_time(timestamp):
    """마지막 처리 시간 저장"""
    state = {'last_processed_time': timestamp}
    with open(STATE_FILE, 'w', encoding='utf-8') as f:
        json.dump(state, f, ensure_ascii=False, indent=2)

def get_new_emails(gmail_service):
    """신규 메일 조회 (이전 처리 시간 이후의 메일만)"""
    past_date = (datetime.now() - timedelta(days=7)).strftime('%Y/%m/%d')
    last_time = get_last_processed_time()

    query = f'after:{past_date}'
    if last_time:
        query += f' after:{last_time}'

    try:
        results = gmail_service.users().messages().list(userId='me', q=query, maxResults=50).execute()
        messages = results.get('messages', [])

        emails = []
        for message in messages:
            msg = gmail_service.users().messages().get(userId='me', id=message['id'], format='full').execute()
            headers = msg['payload']['headers']

            subject = next((h['value'] for h in headers if h['name'] == 'Subject'), '제목 없음')
            sender = next((h['value'] for h in headers if h['name'] == 'From'), '발신자 미확인')
            date = next((h['value'] for h in headers if h['name'] == 'Date'), '')

            # 보낸 사람 이름만 추출
            sender_name = sender.split('<')[0].strip() if '<' in sender else sender

            # 본문 추출
            body = ''
            if 'parts' in msg['payload']:
                for part in msg['payload']['parts']:
                    if part['mimeType'] == 'text/plain':
                        data = part['body'].get('data', '')
                        if data:
                            body = base64.urlsafe_b64decode(data).decode('utf-8')
                            break
            else:
                data = msg['payload']['body'].get('data', '')
                if data:
                    body = base64.urlsafe_b64decode(data).decode('utf-8')

            # 본문에서 불필요한 부분 제거 (이전 메일 포함 제외)
            if '\n\n---' in body:
                body = body.split('\n\n---')[0]
            if 'On ' in body and 'wrote:' in body:
                lines = body.split('\n')
                body = '\n'.join([l for l in lines if not l.startswith('>')])

            emails.append({
                'subject': subject,
                'sender_name': sender_name,
                'date': date,
                'body': body[:500] if body else '본문 없음'
            })

        return emails

    except Exception as e:
        print(f"이메일 조회 오류: {e}")
        return []

def create_word_report(emails):
    """Word 문서 생성"""
    doc = Document()

    # 제목
    title = doc.add_heading('주간 메일 요약', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 기간
    period = doc.add_paragraph()
    period.add_run('생성일: ').bold = True
    period.add_run(datetime.now().strftime("%Y년 %m월 %d일 %H:%M:%S"))
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 요약 통계
    doc.add_heading(f'총 {len(emails)}개 메일 요약', level=1)
    doc.add_paragraph()

    # 메일 상세
    for i, email in enumerate(emails, 1):
        doc.add_heading(f"{i}. {email['subject']}", level=2)
        doc.add_paragraph(f"발신자: {email['sender_name']}")
        doc.add_paragraph(f"날짜: {email['date']}")
        doc.add_paragraph(f"내용: {email['body']}")
        doc.add_paragraph()

    return doc

def upload_to_drive(drive_service, doc, folder_id):
    """Google Drive에 파일 업로드"""
    try:
        # Word 파일을 메모리에 저장
        file_name = f"Weekly_Mail_Summary_{datetime.now().strftime('%Y_%m_%d')}.docx"
        file_buffer = io.BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)

        file_metadata = {
            'name': file_name,
            'parents': [folder_id]
        }

        media = MediaIoBaseUpload(file_buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, webViewLink'
        ).execute()

        print(f"Google Drive 업로드 완료: {file['webViewLink']}")
        return file['webViewLink']

    except Exception as e:
        print(f"Drive 업로드 오류: {e}")
        return None

def create_calendar_event(calendar_service, doc_link, email_count):
    """Google Calendar에 이벤트 추가"""
    try:
        event = {
            'summary': f'주간 메일 요약 ({email_count}개)',
            'description': f'주간 메일 요약 보고서\n문서 링크: {doc_link}',
            'start': {
                'dateTime': datetime.now().isoformat(),
                'timeZone': 'Asia/Seoul',
            },
            'end': {
                'dateTime': (datetime.now() + timedelta(hours=1)).isoformat(),
                'timeZone': 'Asia/Seoul',
            },
            'visibility': 'private',
        }

        event = calendar_service.events().insert(calendarId='primary', body=event).execute()
        print(f"Calendar 이벤트 생성 완료")
        return True

    except Exception as e:
        print(f"Calendar 이벤트 생성 오류: {e}")
        return False

def main():
    try:
        print("=== 주간 메일 요약 루틴 시작 ===")
        print(f"실행 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # 환경 변수에서 인증 정보 설정
        setup_credentials()

        # 서비스 인증
        gmail, calendar, drive, creds = get_services()

        # 신규 메일 조회
        print("신규 메일 조회 중...")
        emails = get_new_emails(gmail)

        if not emails:
            print("처리할 신규 메일이 없습니다")
            return

        print(f"신규 메일 {len(emails)}개 발견")

        # Word 문서 생성
        print("Word 문서 생성 중...")
        doc = create_word_report(emails)

        # Google Drive 업로드
        print("Google Drive 업로드 중...")
        doc_link = upload_to_drive(drive, doc, MAIL_FOLDER_ID)

        if not doc_link:
            print("Drive 업로드 실패")
            return

        # Google Calendar 등록
        print("Google Calendar 등록 중...")
        create_calendar_event(calendar, doc_link, len(emails))

        # 마지막 처리 시간 저장
        save_last_processed_time(datetime.now().isoformat())

        print("=== 루틴 완료 ===")

    except Exception as e:
        print(f"오류 발생: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()
